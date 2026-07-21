from __future__ import annotations

import csv
import io
import json
from collections.abc import Iterable
from urllib.parse import urlsplit

from bs4 import BeautifulSoup

from .constants import (
    _CAPTCHA_PATTERNS,
    _MAX_EXTRACTED_CHARS,
    _MAX_PDF_PAGES,
    _MAX_XLSX_ROWS,
    _MIN_ARTICLE_WORDS,
    _PAYWALL_PATTERNS,
)
from .models import ContentBlock, _ExtractionResult, _ParseError


def _cap_text(text: str) -> str:
    text = text.strip()
    if len(text) <= _MAX_EXTRACTED_CHARS:
        return text
    return text[:_MAX_EXTRACTED_CHARS].rstrip()


def _cap_blocks(blocks: list[ContentBlock]) -> list[ContentBlock]:
    capped: list[ContentBlock] = []
    current_len = 0
    for block in blocks:
        block_len = len(block.text) + 1
        if current_len + block_len > _MAX_EXTRACTED_CHARS:
            if not capped:
                capped.append(
                    ContentBlock(type=block.type, text=block.text[:_MAX_EXTRACTED_CHARS].rstrip())
                )
            break
        capped.append(block)
        current_len += block_len
    return capped


def _cap_extraction_result(result: _ExtractionResult) -> _ExtractionResult:
    if len(result.text) <= _MAX_EXTRACTED_CHARS:
        return result
    blocks = _cap_blocks(result.blocks) if result.blocks else []
    return _ExtractionResult(
        title=result.title,
        text=_cap_text(result.text),
        blocks=blocks,
        published_at=result.published_at,
    )


def _trafilatura_xml_to_blocks(xml_str: str) -> list[ContentBlock]:
    blocks: list[ContentBlock] = []
    if not xml_str:
        return blocks
    try:
        soup = BeautifulSoup(xml_str, "lxml-xml")
        for elem in soup.find_all(True):
            tag = elem.name
            text = elem.get_text(" ", strip=True)
            if not text:
                continue
            if tag == "head":
                blocks.append(ContentBlock(type="heading", text=text))
            elif tag == "p":
                blocks.append(ContentBlock(type="paragraph", text=text))
            elif tag == "item":
                blocks.append(ContentBlock(type="list_item", text=text))
            elif tag == "quote":
                blocks.append(ContentBlock(type="blockquote", text=text))
            elif tag == "code":
                blocks.append(ContentBlock(type="code", text=text))
            elif tag == "table":
                rows: list[str] = []
                for row in elem.find_all("row"):
                    cells = [cell.get_text(" ", strip=True) for cell in row.find_all("cell")]
                    rows.append(" | ".join(cells))
                blocks.append(ContentBlock(type="table", text="\n".join(rows)))
            elif tag == "figcaption":
                blocks.append(ContentBlock(type="caption", text=text))
    except Exception:
        pass
    return blocks


def _soup_to_blocks(soup: BeautifulSoup) -> list[ContentBlock]:
    blocks: list[ContentBlock] = []
    for node in soup.find_all(
        ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "blockquote", "pre", "code", "table", "figcaption"]
    ):
        tag_name = node.name
        text = node.get_text(" ", strip=True)
        if not text:
            continue
        if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            blocks.append(ContentBlock(type="heading", text=text))
        elif tag_name == "li":
            blocks.append(ContentBlock(type="list_item", text=text))
        elif tag_name == "blockquote":
            blocks.append(ContentBlock(type="blockquote", text=text))
        elif tag_name in ("pre", "code"):
            blocks.append(ContentBlock(type="code", text=text))
        elif tag_name == "table":
            rows: list[str] = []
            for tr in node.find_all("tr"):
                cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
                rows.append(" | ".join(cells))
            blocks.append(ContentBlock(type="table", text="\n".join(rows)))
        elif tag_name == "figcaption":
            blocks.append(ContentBlock(type="caption", text=text))
        else:
            blocks.append(ContentBlock(type="paragraph", text=text))
    return blocks


def _extract_with_trafilatura(html: str) -> _ExtractionResult | None:
    try:
        import trafilatura

        json_result = trafilatura.extract(
            html,
            output_format="json",
            with_metadata=True,
            include_comments=False,
            include_tables=True,
            include_images=False,
        )
        if not json_result:
            return None
        data = json.loads(json_result)
        title = (data.get("title") or "").strip()
        text = (data.get("text") or "").strip()
        if len(text.split()) < _MIN_ARTICLE_WORDS:
            return None
        published_at = (data.get("date") or "").strip() or None

        xml_result = trafilatura.extract(
            html,
            output_format="xml",
            include_comments=False,
            include_tables=True,
            include_images=False,
        )
        blocks = _trafilatura_xml_to_blocks(xml_result or "")
        if not blocks:
            for line in text.splitlines():
                line = line.strip()
                if line:
                    blocks.append(ContentBlock(type="paragraph", text=line))
        return _cap_extraction_result(
            _ExtractionResult(title=title, text=text, blocks=blocks, published_at=published_at)
        )
    except Exception:
        return None


def _extract_with_readability(html: str) -> _ExtractionResult | None:
    try:
        from readability import Document

        doc = Document(html)
        title = (doc.title() or "").strip()
        summary = doc.summary()
        if not summary or len(summary) < 100:
            return None
        soup = BeautifulSoup(summary, "lxml")
        blocks = _soup_to_blocks(soup)
        text = "\n".join(b.text for b in blocks).strip()
        if len(text.split()) < _MIN_ARTICLE_WORDS:
            return None
        return _cap_extraction_result(_ExtractionResult(title=title, text=text, blocks=blocks))
    except Exception:
        return None


def _extract_with_bs4(html: str) -> _ExtractionResult:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup.select("script,style,noscript,header,footer,nav,aside,form,iframe,svg"):
        tag.decompose()

    title = _extract_title_from_soup(soup)
    section = (
        soup.find("article")
        or soup.find("main")
        or soup.select_one("[role='main']")
        or soup.find("body")
        or soup
    )

    blocks = _soup_to_blocks(section)
    total_words = len(" ".join(b.text for b in blocks).split())
    if total_words < 120:
        blocks = []
        for line in _clean_blocks(line for line in section.get_text("\n").splitlines()):
            blocks.append(ContentBlock(type="paragraph", text=line))

    text = "\n".join(b.text for b in blocks).strip()
    if len(text) > _MAX_EXTRACTED_CHARS:
        truncated = _cap_blocks(blocks)
        text = "\n".join(b.text for b in truncated).strip()
        blocks = truncated

    return _ExtractionResult(title=title, text=text, blocks=blocks)


def _extract_html(html: str) -> _ExtractionResult:
    result = _extract_with_trafilatura(html)
    if result is not None:
        return result
    result = _extract_with_readability(html)
    if result is not None:
        return result
    return _extract_with_bs4(html)


def _parse_pdf(data: bytes) -> _ExtractionResult:
    try:
        import fitz

        doc = fitz.open(stream=data, filetype="pdf")
        try:
            if doc.page_count > _MAX_PDF_PAGES:
                raise _ParseError("non_text_content", f"PDF exceeds {_MAX_PDF_PAGES} pages")
            title = (doc.metadata.get("title") or "").strip()
            if not title and len(doc) > 0:
                blocks = doc[0].get_text("blocks")
                for block in blocks:
                    block_text = block[4].strip()
                    if block_text and 10 < len(block_text) < 200:
                        title = block_text
                        break
            full_text: list[str] = []
            for page in doc:
                full_text.append(page.get_text())
            text = _cap_text("\n".join(full_text))
            if not text:
                raise _ParseError("pdf_parse_failed", "Empty PDF text")
            return _ExtractionResult(title=title, text=text)
        finally:
            doc.close()
    except _ParseError:
        raise
    except Exception as exc:
        raise _ParseError("pdf_parse_failed", str(exc)) from exc


def _parse_docx(data: bytes) -> _ExtractionResult:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        title = (doc.core_properties.title or "").strip()
        paragraphs: list[str] = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                paragraphs.append(txt)
        if not title and paragraphs:
            for p in paragraphs:
                if len(p) > 5:
                    title = p
                    break
        text = _cap_text("\n".join(paragraphs))
        if not text:
            raise _ParseError("non_text_content", "Empty DOCX")
        return _ExtractionResult(title=title, text=text)
    except _ParseError:
        raise
    except Exception as exc:
        raise _ParseError("non_text_content", str(exc)) from exc


def _parse_xlsx(data: bytes) -> _ExtractionResult:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        try:
            rows: list[str] = []
            row_count = 0
            for sheet in wb:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        rows.append(" | ".join(cells))
                    row_count += 1
                    if row_count >= _MAX_XLSX_ROWS:
                        break
                if row_count >= _MAX_XLSX_ROWS:
                    break
            text = _cap_text("\n".join(rows))
            if not text:
                raise _ParseError("non_text_content", "Empty XLSX")
            title = wb.properties.title or ""
        finally:
            wb.close()
        return _ExtractionResult(title=title, text=text)
    except _ParseError:
        raise
    except Exception as exc:
        raise _ParseError("non_text_content", str(exc)) from exc


def _parse_csv(data: bytes) -> _ExtractionResult:
    try:
        text = data.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        rows: list[str] = []
        for row in reader:
            if row:
                rows.append(" | ".join(row))
        text = _cap_text("\n".join(rows))
        if not text:
            raise _ParseError("non_text_content", "Empty CSV")
        return _ExtractionResult(title="", text=text)
    except _ParseError:
        raise
    except Exception as exc:
        raise _ParseError("non_text_content", str(exc)) from exc


def _parse_plaintext(data: bytes, content_type: str | None = None) -> _ExtractionResult:
    text = _cap_text(data.decode("utf-8", errors="ignore"))
    if not text:
        raise _ParseError("non_text_content", "Empty text file")
    return _ExtractionResult(title="", text=text)


def _extract_published_at(html: str) -> str | None:
    soup = BeautifulSoup(html, "lxml")

    og_time = soup.find("meta", attrs={"property": "article:published_time"})
    if og_time and og_time.get("content"):
        return str(og_time["content"]).strip()

    for script in soup.select("script[type='application/ld+json']"):
        try:
            data = json.loads(script.string or "")
        except (json.JSONDecodeError, TypeError):
            continue
        if isinstance(data, dict):
            for key in ("datePublished", "dateCreated", "dateModified", "uploadDate"):
                pub = data.get(key)
                if pub:
                    return str(pub).strip()
        if isinstance(data, list):
            for item in data:
                if isinstance(item, dict):
                    for key in ("datePublished", "dateCreated", "dateModified", "uploadDate"):
                        pub = item.get(key)
                        if pub:
                            return str(pub).strip()

    meta_names = (
        "date",
        "pubdate",
        "dc.date",
        "dc.date.issued",
        "prism.publicationDate",
        "sailthru.date",
        "parsely-pub-date",
        "datePublished",
        "dateModified",
        "uploadDate",
    )
    for meta_name in meta_names:
        tag = soup.find("meta", attrs={"name": meta_name})
        if tag and tag.get("content"):
            return str(tag["content"]).strip()

    time_tag = soup.find("time", attrs={"datetime": True})
    if time_tag:
        return str(time_tag["datetime"]).strip()

    return None


def _detect_fetch_error(html: str | None, status: int | None, text: str) -> str | None:
    if html:
        if len(html) > 8000 and len(text.split()) < 40:
            return "empty_js_shell"

        weak_text = len(text.split()) < 80
        html_lower = html.lower()

        captcha_matched = weak_text and any(p.search(html_lower) for p in _CAPTCHA_PATTERNS)
        if captcha_matched:
            return "captcha"

        if weak_text:
            paywall_matched = any(p.search(html_lower) for p in _PAYWALL_PATTERNS)
            if paywall_matched:
                return "paywall"

    return None


def _error_for_status(status: int | None) -> str:
    if status == 403:
        return "blocked_403"
    if status == 401:
        return "login_required"
    if status == 404:
        return "scrape_failed"
    if status == 410:
        return "scrape_failed"
    if status is not None and status >= 500:
        return "scrape_failed"
    return "scrape_failed"


def _guess_format(url: str, content_type: str) -> str:
    ct = content_type.lower()
    if "html" in ct or "xhtml" in ct:
        return "html"
    if "pdf" in ct:
        return "pdf"
    if "wordprocessingml.document" in ct or "msword" in ct:
        return "docx"
    if "spreadsheetml.sheet" in ct or "excel" in ct:
        return "xlsx"
    if "csv" in ct:
        return "csv"
    if "text/plain" in ct:
        return "txt"
    if "text/markdown" in ct or "text/x-markdown" in ct:
        return "md"

    path = urlsplit(url).path.lower()
    if path.endswith(".pdf"):
        return "pdf"
    if path.endswith(".docx"):
        return "docx"
    if path.endswith(".xlsx"):
        return "xlsx"
    if path.endswith(".csv"):
        return "csv"
    if path.endswith(".txt"):
        return "txt"
    if path.endswith(".md"):
        return "md"

    if ct.startswith("text/"):
        return "html"

    if not ct:
        return "html"

    return ""


def _extract_title_from_soup(soup: BeautifulSoup) -> str:
    og_title = soup.find("meta", attrs={"property": "og:title"})
    if og_title and og_title.get("content"):
        return str(og_title["content"]).strip()
    title_tag = soup.find("title")
    if title_tag and title_tag.text:
        return title_tag.text.strip()
    h1 = soup.find("h1")
    if h1 and h1.text:
        return h1.text.strip()
    return ""


def _clean_blocks(lines: Iterable[str]) -> Iterable[str]:
    seen: set[str] = set()
    for raw in lines:
        line = " ".join(raw.split()).strip()
        if len(line) < 35:
            continue
        lowered = line.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        yield line
