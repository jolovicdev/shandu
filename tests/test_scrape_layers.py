from __future__ import annotations

import asyncio
import io

import aiohttp
import fitz
from docx import Document
from openpyxl import Workbook

from shandu.services.scrape import ScrapeService
from shandu.services.scrape.extraction import (
    _detect_fetch_error,
    _extract_html,
    _extract_published_at,
    _extract_with_bs4,
    _extract_with_readability,
    _extract_with_trafilatura,
    _parse_csv,
    _parse_docx,
    _parse_pdf,
    _parse_plaintext,
    _parse_xlsx,
)


def _long_html(title: str = "Title", words: int = 100) -> str:
    text = " ".join(["word"] * words)
    return f"<html><head><title>{title}</title></head><body><article><h1>{title}</h1><p>{text}</p><p>{text}</p></article></body></html>"


# ---------------------------------------------------------------------------
# Article extractors
# ---------------------------------------------------------------------------

def test_trafilatura_extracts_title_text_and_blocks() -> None:
    html = _long_html("Trafilatura Article", words=100)
    result = _extract_with_trafilatura(html)
    assert result is not None
    assert result.title == "Trafilatura Article"
    assert len(result.text.split()) >= 100
    assert len(result.blocks) > 0


def test_readability_extracts_title_text_and_blocks() -> None:
    html = _long_html("Readability Article", words=100)
    result = _extract_with_readability(html)
    assert result is not None
    assert result.title == "Readability Article"
    assert len(result.text.split()) >= 100
    assert len(result.blocks) > 0


def test_bs4_fallback_extracts_title_text_and_blocks() -> None:
    html = _long_html("BS4 Article", words=100)
    result = _extract_with_bs4(html)
    assert result.title == "BS4 Article"
    assert len(result.text.split()) >= 100
    assert len(result.blocks) > 0


def test_extract_cascade_prefers_trafilatura_then_readability_then_bs4() -> None:
    service = ScrapeService()
    title, text = service._extract(_long_html("Cascade", words=120))
    assert title == "Cascade"
    assert len(text.split()) >= 120


def test_trafilatura_returns_none_for_short_content() -> None:
    html = "<html><body><article><p>Short text.</p></article></body></html>"
    assert _extract_with_trafilatura(html) is None


def test_readability_returns_none_for_short_content() -> None:
    html = "<html><body><article><p>Short text.</p></article></body></html>"
    assert _extract_with_readability(html) is None


# ---------------------------------------------------------------------------
# Document parsers
# ---------------------------------------------------------------------------

def test_parse_pdf_extracts_title_and_text() -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PDF Document Title")
    page.insert_text((72, 100), "This is the body text of the PDF document.")
    data = doc.tobytes()
    result = _parse_pdf(data)
    assert result.title == "PDF Document Title"
    assert "body text" in result.text


def test_parse_docx_extracts_title_and_text() -> None:
    document = Document()
    document.add_heading("DOCX Title", 0)
    document.add_paragraph("This is a paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    result = _parse_docx(buffer.read())
    assert result.title == "DOCX Title"
    assert "paragraph" in result.text


def test_parse_xlsx_extracts_rows() -> None:
    wb = Workbook()
    ws = wb.active
    ws.append(["Name", "Value"])
    ws.append(["Alice", 100])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    result = _parse_xlsx(buffer.read())
    assert "Alice | 100" in result.text


def test_parse_csv_extracts_rows() -> None:
    result = _parse_csv(b"Name,Value\nAlice,100\nBob,200")
    assert "Alice | 100" in result.text
    assert "Bob | 200" in result.text


def test_parse_plaintext_extracts_text() -> None:
    result = _parse_plaintext(b"Hello world", "text/plain")
    assert result.text == "Hello world"


# ---------------------------------------------------------------------------
# Date extraction
# ---------------------------------------------------------------------------

def test_extract_published_at_finds_new_meta_names() -> None:
    html = (
        '<html><head>'
        '<meta name="prism.publicationDate" content="2024-03-01">'
        '<meta name="sailthru.date" content="2024-03-02">'
        '</head><body></body></html>'
    )
    assert _extract_published_at(html) == "2024-03-01"


def test_extract_published_at_finds_dc_date_issued() -> None:
    html = '<html><head><meta name="dc.date.issued" content="2024-04-01"></head><body></body></html>'
    assert _extract_published_at(html) == "2024-04-01"


def test_extract_published_at_finds_parsely_pub_date() -> None:
    html = '<html><head><meta name="parsely-pub-date" content="2024-05-01"></head><body></body></html>'
    assert _extract_published_at(html) == "2024-05-01"


def test_extract_published_at_finds_upload_date() -> None:
    html = (
        '<html><head>'
        '<script type="application/ld+json">{"uploadDate":"2024-06-01"}</script>'
        '</head><body></body></html>'
    )
    assert _extract_published_at(html) == "2024-06-01"


# ---------------------------------------------------------------------------
# Paywall / blocked detection
# ---------------------------------------------------------------------------

def test_detect_fetch_error_detects_paywall() -> None:
    html = '<html><body><div class="paywall">Subscribe to read more.</div></body></html>'
    assert _detect_fetch_error(html, 200, "") == "paywall"


def test_detect_fetch_error_detects_captcha() -> None:
    html = '<html><body><div class="g-recaptcha"></div></body></html>'
    assert _detect_fetch_error(html, 200, "") == "captcha"


def test_detect_fetch_error_detects_empty_js_shell() -> None:
    html = '<html><body>' + ' ' * 9000 + '<div id="root"></div></body></html>'
    assert _detect_fetch_error(html, 200, "") == "empty_js_shell"


def test_detect_fetch_error_returns_none_for_normal_page() -> None:
    html = '<html><body><p>This is a normal page with plenty of content.</p></body></html>'
    assert _detect_fetch_error(html, 200, "normal content here") is None


def test_detect_fetch_error_ignores_captcha_marker_when_text_is_strong() -> None:
    html = '<html><body><article><p>' + " ".join(["word"] * 120) + '</p></article><script>g-recaptcha</script></body></html>'
    assert _detect_fetch_error(html, 200, " ".join(["word"] * 120)) is None


def test_extract_html_caps_long_successful_extraction() -> None:
    html = _long_html("Long Article", words=5000)
    result = _extract_html(html)
    assert len(result.text) <= 18000


# ---------------------------------------------------------------------------
# Retry policy
# ---------------------------------------------------------------------------

def test_scrape_retries_on_429_and_eventually_succeeds() -> None:
    service = ScrapeService()
    call_count = 0

    class FakeSession:
        closed = False
        async def close(self):
            pass
        def get(self, *args, **kwargs):
            nonlocal call_count
            call_count += 1
            class FakeResponse:
                url = "https://example.com"
                headers = {"content-type": "text/html"}
                status = 429 if call_count <= 2 else 200
                async def __aenter__(self):
                    return self
                async def __aexit__(self, *args):
                    pass
                def raise_for_status(self):
                    pass
                async def text(self, errors=None):
                    return _long_html(words=120)
            return FakeResponse()

    fake = FakeSession()
    result = asyncio.run(service.scrape("https://example.com", session=fake))
    assert result.fetch_error is None
    assert call_count == 3
    assert service._retry_count == 2


def test_scrape_retries_on_transient_client_error() -> None:
    service = ScrapeService()
    service._backoff_delay = lambda attempt: 0
    call_count = 0

    class FakeSession:
        closed = False
        async def close(self):
            pass
        def get(self, *args, **kwargs):
            nonlocal call_count
            call_count += 1
            if call_count == 1:
                raise aiohttp.ClientConnectionError("connection reset")
            class FakeResponse:
                url = "https://example.com"
                headers = {"content-type": "text/html"}
                status = 200
                async def __aenter__(self):
                    return self
                async def __aexit__(self, *args):
                    pass
                async def text(self, errors=None):
                    return _long_html(words=120)
            return FakeResponse()

    fake = FakeSession()
    result = asyncio.run(service.scrape("https://example.com", session=fake))
    assert result.fetch_error is None
    assert call_count == 2
    assert service._retry_count == 1


def test_scrape_does_not_retry_404() -> None:
    service = ScrapeService()
    call_count = 0

    class FakeSession:
        closed = False
        async def close(self):
            pass
        def get(self, *args, **kwargs):
            nonlocal call_count
            call_count += 1
            class FakeResponse:
                url = "https://example.com"
                headers = {"content-type": "text/html"}
                status = 404
                async def __aenter__(self):
                    return self
                async def __aexit__(self, *args):
                    pass
                def raise_for_status(self):
                    pass
                async def text(self, errors=None):
                    return "not found"
            return FakeResponse()

    fake = FakeSession()
    result = asyncio.run(service.scrape("https://example.com", session=fake))
    assert result.fetch_error == "scrape_failed"
    assert call_count == 1
    assert service._retry_count == 0


# ---------------------------------------------------------------------------
# Content type / format detection
# ---------------------------------------------------------------------------

def test_scrape_service_detects_pdf_from_content_type() -> None:
    service = ScrapeService()

    class _FakeContent:
        async def read(self) -> bytes:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "PDF Title")
            return doc.tobytes()

    class FakeSession:
        closed = False
        async def close(self):
            pass
        def get(self, *args, **kwargs):
            class FakeResponse:
                url = "https://example.com/doc.pdf"
                headers = {"content-type": "application/pdf"}
                status = 200
                content = _FakeContent()
                async def __aenter__(self):
                    return self
                async def __aexit__(self, *args):
                    pass
                def raise_for_status(self):
                    pass
            return FakeResponse()

    fake = FakeSession()
    result = asyncio.run(service.scrape("https://example.com/doc.pdf", session=fake))
    assert result.fetch_error is None
    assert result.content_type == "application/pdf"
    assert "PDF Title" in result.text


def test_scrape_service_detects_docx_from_url_extension() -> None:
    service = ScrapeService()

    class _FakeContent:
        async def read(self) -> bytes:
            document = Document()
            document.add_heading("DOCX Title", 0)
            document.add_paragraph("Hello from docx.")
            buffer = io.BytesIO()
            document.save(buffer)
            return buffer.getvalue()

    class FakeSession:
        closed = False
        async def close(self):
            pass
        def get(self, *args, **kwargs):
            class FakeResponse:
                url = "https://example.com/doc.docx"
                headers = {"content-type": "application/octet-stream"}
                status = 200
                content = _FakeContent()
                async def __aenter__(self):
                    return self
                async def __aexit__(self, *args):
                    pass
                def raise_for_status(self):
                    pass
            return FakeResponse()

    fake = FakeSession()
    result = asyncio.run(service.scrape("https://example.com/doc.docx", session=fake))
    assert result.fetch_error is None
    assert "DOCX Title" in result.text


def test_scrape_service_rejects_oversized_html_before_decoding() -> None:
    service = ScrapeService()

    class FakeSession:
        closed = False
        async def close(self):
            pass
        def get(self, *args, **kwargs):
            class FakeResponse:
                url = "https://example.com"
                headers = {"content-type": "text/html", "content-length": str(11 * 1024 * 1024)}
                status = 200
                async def __aenter__(self):
                    return self
                async def __aexit__(self, *args):
                    pass
                async def text(self, errors=None):
                    raise AssertionError("oversized HTML should not be decoded")
            return FakeResponse()

    fake = FakeSession()
    result = asyncio.run(service.scrape("https://example.com", session=fake))
    assert result.fetch_error == "non_text_content"


# ---------------------------------------------------------------------------
# Structured blocks
# ---------------------------------------------------------------------------

def test_extract_produces_structured_blocks() -> None:
    long_para = " ".join(["word"] * 60)
    html = (
        f"<html><body><article>"
        f"<h2>Heading</h2>"
        f"<p>{long_para}</p>"
        f"<p>{long_para}</p>"
        f"<ul><li>List item one with enough words.</li><li>List item two with enough words.</li></ul>"
        f"<blockquote>A quote that is long enough to pass the filter.</blockquote>"
        f"<pre><code>print(hello)</code></pre>"
        f"</article></body></html>"
    )
    result = _extract_with_bs4(html)
    types = [b.type for b in result.blocks]
    assert "heading" in types
    assert "paragraph" in types
    assert "list_item" in types
    assert "blockquote" in types
    assert "code" in types


def test_scraped_page_defaults_blocks_to_empty_list() -> None:
    from shandu.services.scrape import ScrapedPage
    page = ScrapedPage(
        requested_url="https://example.com",
        url="https://example.com",
        title="Test",
        text="text",
        domain="example.com",
    )
    assert page.blocks == []
